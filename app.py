from dotenv import load_dotenv
from flask_sqlalchemy import SQLAlchemy
import requests
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from flask import Flask, request, send_file, jsonify, send_from_directory
from flask_cors import CORS
from docx import Document
from io import BytesIO
from pathlib import Path
import tempfile
from flask import make_response
from docx.shared import Pt, RGBColor
import traceback
import os
from datetime import datetime, timezone
from docx2pdf import convert


load_dotenv()  # загружает .env


app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://gen_user:12345678i@109.73.193.193:5432/default_db'
app.config['SECRET_KEY'] = 'your-secret-key-here'
db = SQLAlchemy(app)
login_manager = LoginManager(app)
CORS(app, 
    resources={r"/*": {"origins": "http://127.0.0.1:5500"}},
    supports_credentials=True  # Разрешаем передачу кук
)

DEFAULT_FONT_NAME = "Times New Roman"
DEFAULT_FONT_SIZE = 12
DEFAULT_FONT_COLOR = RGBColor(0, 0, 0)

# Модели
class User(UserMixin, db.Model):
    __tablename__ = 'users'
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True)
    email = db.Column(db.String(100), unique=True)
    password_hash = db.Column(db.String(512))
    documents = db.relationship('DocumentHistory', backref='user', lazy=True)

class DocumentHistory(db.Model):
    __tablename__ = 'document_history'
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('users.id'))
    template_name = db.Column(db.String(100))
    generated_at = db.Column(db.DateTime, default=datetime.utcnow)
    download_link = db.Column(db.String(255))  # Можно оставить для обратной совместимости
    document_name = db.Column(db.String(100))
    file_data = db.Column(db.LargeBinary)  # Бинарные данные файла
    file_type = db.Column(db.String(50))   # Тип файла: 'docx' или 'pdf'

def apply_styles_to_paragraph(paragraph):
    for run in paragraph.runs:
        font = run.font
        font.name = DEFAULT_FONT_NAME
        font.size = Pt(DEFAULT_FONT_SIZE)
        font.color.rgb = DEFAULT_FONT_COLOR

@app.route("/")
def home():
    return send_from_directory("static", "create_document.html")

@app.route("/<path:filename>")
def static_files(filename):
    return send_from_directory("static", filename)

# Новый обработчик для POST запросов без имени файла в URL
@app.route('/generate_d', methods=['POST'])
def generate_docx_document():
    try:
        data = request.json
        name_of_doc = data.get('filename')
        if not name_of_doc:
            return jsonify({"error": "Filename not provided"}), 400

        # Удаляем параметры из данных, чтобы не пытаться их подставить
        form_data = {k: v for k, v in data.items() if k != 'filename'}

        template_path = os.path.join('templates', name_of_doc)
        if not os.path.exists(template_path):
            return jsonify({"error": f"Template file not found: {template_path}"}), 404

        doc = Document(template_path)

        for paragraph in doc.paragraphs:
            for key, value in form_data.items():
                placeholder = f'{{{{ {key} }}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    apply_styles_to_paragraph(paragraph)

        output_path = 'generated_document.docx'
        doc.save(output_path)

        return send_file(output_path, as_attachment=True, download_name='generated_document.docx')

    except Exception as e:
        app.logger.error(f"Ошибка: {e}")
        app.logger.error(traceback.format_exc())
        return jsonify({"error": "Internal Server Error", "details": str(e)}), 500

def generate_document_name(template_name, data):
    date_str = datetime.now().strftime("%d.%m.%Y")
    if template_name == "purchase&sale_agreement.docx":
        return f"ДКП {data.get('name_of_seller', '')} - {data.get('name_of_buyer', '')} {date_str}"
    else:
        return f"{template_name} {date_str}"

# Старый обработчик (оставлен для совместимости)
@app.route('/generate_d/<name_of_doc>', methods=['POST'])
def generate_docx_document_with_path(name_of_doc):
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        template_path = os.path.join('templates', name_of_doc)
        if not os.path.exists(template_path):
            return jsonify({"error": f"Template file not found: {template_path}"}), 404

        doc = Document(template_path)

        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f'{{{{ {key} }}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    apply_styles_to_paragraph(paragraph)

        # Сохраняем документ в буфер памяти
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        file_data = buffer.read()

        doc_history = DocumentHistory(
            user_id=current_user.id,
            document_name=generate_document_name(name_of_doc, data),
            template_name=name_of_doc,
            download_link=f"/download_doc/{name_of_doc}",
            file_data=file_data,
            file_type='docx'
        )
        db.session.add(doc_history)
        db.session.commit()

        # Возвращаем файл для скачивания
        buffer.seek(0)  # Сбрасываем позицию буфера
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{doc_history.document_name}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        app.logger.error(f"Ошибка: {e}")
        app.logger.error(traceback.format_exc())
        return jsonify({"error": "Internal Server Error", "details": str(e)}), 500
    
@app.route('/generate_p/<name_of_doc>', methods=['POST'])
def generate_pdf_document(name_of_doc):
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        template_path = os.path.join('templates', name_of_doc)
        if not os.path.exists(template_path):
            return jsonify({"error": f"Template file not found: {template_path}"}), 404

        # Обработка DOCX шаблона
        doc = Document(template_path)
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                placeholder = f'{{{{ {key} }}}}'
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    apply_styles_to_paragraph(paragraph)

        # Создаем временную директорию
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx_path = os.path.join(temp_dir, 'temp.docx')
            output_pdf_path = os.path.join(temp_dir, 'output.pdf')

            # Сохраняем DOCX во временный файл
            doc.save(temp_docx_path)

            # Конвертируем в PDF
            convert(temp_docx_path, output_pdf_path)

            # Читаем PDF в память
            with open(output_pdf_path, 'rb') as f:
                pdf_data = f.read()

            # Сохраняем в базу данных
            doc_history = DocumentHistory(
                user_id=current_user.id,
                document_name=generate_document_name(name_of_doc, data),
                template_name=name_of_doc,
                download_link=f"/download_doc/{name_of_doc}",
                file_data=pdf_data,
                file_type='pdf'
            )
            db.session.add(doc_history)
            db.session.commit()

            # Возвращаем файл для предпросмотра
            return send_file(
                BytesIO(pdf_data),
                mimetype='application/pdf',
                as_attachment=False,
                download_name=f"{doc_history.document_name}.pdf"
            )

    except Exception as e:
        app.logger.error(f"Ошибка: {e}")
        app.logger.error(traceback.format_exc())
        return jsonify({"error": "Internal Server Error", "details": str(e)}), 500
    

@app.route('/download_doc/<int:doc_id>')
@login_required
def download_document(doc_id):
    doc = DocumentHistory.query.filter_by(
        id=doc_id,
        user_id=current_user.id
    ).first()

    if not doc:
        return jsonify({"error": "Document not found"}), 404

    # Определяем MIME-тип и расширение файла
    if doc.file_type == 'pdf':
        mimetype = 'application/pdf'
        extension = 'pdf'
    else:
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        extension = 'docx'

    # Создаем временный файл в памяти
    from io import BytesIO
    buffer = BytesIO(doc.file_data)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{doc.document_name}.{extension}",
        mimetype=mimetype
    )


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Регистрация
@app.route('/register', methods=['POST'])
def register():
    data = request.json
    hashed_password = generate_password_hash(data['password'])
    new_user = User(username=data['username'], email=data['email'], password_hash=hashed_password)
    db.session.add(new_user)
    db.session.commit()
    return jsonify({"message": "User created"}), 201

# Вход
@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    user = User.query.filter_by(username=data['username']).first()
    
    if user and check_password_hash(user.password_hash, data['password']):
        login_user(user, remember=True)  # Добавляем remember=True
        return jsonify({
            "message": "Logged in successfully",
            "user": user.username,
            "id": user.id
        }), 200
    return jsonify({"error": "Invalid credentials"}), 401

# Выход
@app.route('/logout', methods=['GET', 'POST'])  # Добавляем POST
@login_required
def logout():
    logout_user()
    return jsonify({"message": "Logged out"})

@app.route('/api/history', methods=['GET'])
@login_required
def get_history():
    try:
        history = DocumentHistory.query.filter_by(
            user_id=current_user.id
        ).order_by(
            DocumentHistory.generated_at.desc()
        ).all()
        
        return jsonify([{
            'id': item.id,
            'template_name': item.template_name,
            'generated_at': item.generated_at.replace(tzinfo=timezone.utc)
                             .astimezone(tz=None)
                             .strftime('%d.%m.%Y %H:%M'),
            'download_link': f"/download_doc/{item.id}",  # Используем новый эндпоинт
            'document_name': item.document_name,
            'file_type': item.file_type
        } for item in history])
    
    except Exception as e:
        return jsonify({"error": str(e)}), 500
    

@app.route('/api/history/<int:doc_id>', methods=['DELETE'])
@login_required
def delete_history_item(doc_id):
    item = DocumentHistory.query.filter_by(id=doc_id, user_id=current_user.id).first()
    if not item:
        return jsonify({"error": "Document not found"}), 404
    
    db.session.delete(item)
    db.session.commit()
    return jsonify({"message": "Deleted"}), 200


@app.route('/api/check-auth')
def check_auth():
    if current_user.is_authenticated:
        return jsonify({
            'authenticated': True,
            'user': {
                'id': current_user.id,
                'username': current_user.username,
                'email': current_user.email
            }
        })
    return jsonify({'authenticated': False})


#чат
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

@app.route("/api/chat", methods=["POST"])
def chat():
    user_message = request.json.get("message")

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "HTTP-Referer": "http://127.0.0.1:5500", 
        "X-Title": "DocumentGenerator"
    }

    payload = {
        "model": "deepseek/deepseek-chat-v3-0324:free",
        "messages": [{"role": "user", "content": user_message}]
    }

    try:
        response = requests.post("https://openrouter.ai/api/v1/chat/completions",
                                 headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        return jsonify({"reply": data["choices"][0]["message"]["content"]})

    except Exception as e:
        print("Ошибка:", e)
        return jsonify({"reply": "Ошибка: не удалось получить ответ от сервера."}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5500, threaded=False)
